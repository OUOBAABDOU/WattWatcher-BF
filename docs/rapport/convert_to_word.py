import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Configuration
ROOT_DIR = Path(__file__).resolve().parents[2]
MD_FILE = ROOT_DIR / "docs" / "rapport" / "rapport_final.md"
WORD_OUTPUT = ROOT_DIR / "docs" / "rapport" / "rapport_final.docx"
CAPTURES_DIR = ROOT_DIR / "docs" / "captures"

def add_heading(doc, text, level):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text):
    """Add a paragraph to the document"""
    if text.strip():
        para = doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return para
    return None

def add_image(doc, image_path, caption=None):
    """Add an image to the document with optional caption"""
    full_path = ROOT_DIR / image_path.replace("../", "")
    if full_path.exists():
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(full_path), width=Inches(6))
            
            if caption:
                caption_para = doc.add_paragraph(caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.runs[0].italic = True
                caption_para.runs[0].font.size = Pt(10)
        except Exception as e:
            doc.add_paragraph(f"[Image non trouvée: {image_path}]")
    else:
        doc.add_paragraph(f"[Image non trouvée: {image_path}]")

def add_table(doc, headers, rows):
    """Add a table to the document"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def parse_markdown_table(table_text):
    """Parse a Markdown table and return headers and rows"""
    lines = table_text.strip().split('\n')
    if len(lines) < 2:
        return [], []
    
    # Parse header
    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    
    # Skip separator line
    rows = []
    for line in lines[2:]:
        if line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    return headers, rows

def convert_markdown_to_docx(md_file, output_file):
    """Convert Markdown file to Word document"""
    doc = Document()
    
    # Set default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    doc.styles['Normal'].font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into lines
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            if level <= 6:
                heading_text = line.lstrip('#').strip()
                add_heading(doc, heading_text, level)
                i += 1
                continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            para = doc.add_paragraph(code_text)
            para.runs[0].font.name = 'Courier New'
            para.runs[0].font.size = Pt(9)
            i += 1
            continue
        
        # Handle images
        if line.strip().startswith('!['):
            # Parse image: ![alt](path)
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                # Check if next line is a caption
                caption = None
                if i + 1 < len(lines) and lines[i + 1].strip().startswith('*'):
                    caption = lines[i + 1].strip().strip('*')
                    i += 1
                add_image(doc, image_path, caption)
                i += 1
                continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            # Collect table lines
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            
            table_text = '\n'.join(table_lines)
            headers, rows = parse_markdown_table(table_text)
            if headers and rows:
                add_table(doc, headers, rows)
            continue
        
        # Handle lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                list_items.append(lines[i].strip()[2:])
                i += 1
            for item in list_items:
                para = doc.add_paragraph(item, style='List Bullet')
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line.strip()):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.', lines[i].strip()):
                list_items.append(lines[i].strip())
                i += 1
            for item in list_items:
                para = doc.add_paragraph(item, style='List Number')
            continue
        
        # Handle empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            para = add_paragraph(doc, line.strip())
            i += 1
            continue
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"Document Word généré : {output_file}")

if __name__ == "__main__":
    print(f"Conversion de {MD_FILE} vers Word...")
    convert_markdown_to_docx(MD_FILE, WORD_OUTPUT)
    print("Conversion terminée avec succès!")
