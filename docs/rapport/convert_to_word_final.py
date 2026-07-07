import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration
ROOT_DIR = Path(__file__).resolve().parents[2]
MD_FILE = ROOT_DIR / "docs" / "rapport" / "rapport_final.md"
WORD_OUTPUT = ROOT_DIR / "docs" / "rapport" / "rapport_final_universitaire.docx"
CAPTURES_DIR = ROOT_DIR / "docs" / "captures"

# Compteurs pour les listes
figure_counter = 0
table_counter = 0
figures_list = []
tables_list = []
annexes_list = []

def clean_inline_text(text):
    """Nettoyer le balisage Markdown simple sans changer le contenu."""
    text = text.replace('**', '').replace('__', '')
    text = text.replace('`', '')
    return text.strip()

def normalize_caption(caption, prefix, number):
    """Uniformiser une légende de figure ou de tableau."""
    caption = clean_inline_text(caption).replace('*', '').strip()
    caption = re.sub(rf'^{prefix}\s*(?:X|\d+)\s*:\s*', '', caption, flags=re.IGNORECASE)
    caption = re.sub(rf'^{prefix}\s+', '', caption, flags=re.IGNORECASE)
    return f"{prefix} {number} : {caption}"

def scan_document_lists(content):
    """Préparer les listes à afficher avant le corps du rapport."""
    scanned_figures = []
    scanned_tables = []
    scanned_annexes = []
    fig_no = 0
    table_no = 0

    for line in content.split('\n'):
        stripped = line.strip()
        if stripped.startswith('*Figure'):
            fig_no += 1
            scanned_figures.append(normalize_caption(stripped.strip('*'), 'Figure', fig_no))
        elif re.match(r'^#{2,6}\s+Tableau', stripped):
            table_no += 1
            title = re.sub(r'^#{2,6}\s+', '', stripped)
            scanned_tables.append(normalize_caption(title, 'Tableau', table_no))
        elif re.match(r'^#{2,6}\s+Annexe\s+[A-Z]\s*:', stripped):
            scanned_annexes.append(clean_inline_text(re.sub(r'^#{2,6}\s+', '', stripped)))

    return scanned_figures, scanned_tables, scanned_annexes

def set_margins(doc, margin_cm=2.5):
    """Définir les marges du document"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

def set_header(doc, text="Rapport de projet tutoré — WattWatcher BF"):
    """Ajouter l'en-tête"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = text
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.runs[0].font.name = 'Times New Roman'
    header_para.runs[0].font.size = Pt(10)

def set_footer(doc):
    """Ajouter le pied de page avec numérotation"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    
    run = footer_para.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

def set_font_style(run, font_name='Times New Roman', size=12, bold=False, italic=False):
    """Définir le style de police"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, space_before=6, space_after=6):
    """Définir le format du paragraphe"""
    para.alignment = alignment
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

def add_heading(doc, text, level):
    """Ajouter un titre avec le style académique"""
    if level == 0:
        heading = doc.add_heading(text, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_style(heading.runs[0], size=18, bold=True)
    elif level == 1:
        heading = doc.add_heading(text, level=2)
        set_font_style(heading.runs[0], size=16, bold=True)
    elif level == 2:
        heading = doc.add_heading(text, level=3)
        set_font_style(heading.runs[0], size=14, bold=True)
    else:
        heading = doc.add_heading(text, level=4)
        set_font_style(heading.runs[0], size=13, bold=True)
    
    set_paragraph_format(heading, space_before=12, space_after=6)
    return heading

def add_paragraph(doc, text):
    """Ajouter un paragraphe avec le style académique"""
    if text.strip():
        para = doc.add_paragraph(clean_inline_text(text))
        set_font_style(para.runs[0], size=12)
        set_paragraph_format(para)
        return para
    return None

def add_image(doc, image_path, caption=None):
    """Ajouter une image avec légende et incrémenter le compteur"""
    global figure_counter, figures_list
    
    full_path = ROOT_DIR / image_path.replace("../", "")
    
    if not full_path.exists():
        captures_dir = ROOT_DIR / "docs" / "captures"
        if captures_dir.exists():
            filename = Path(image_path).name
            for file in captures_dir.glob("*.png"):
                if filename in file.name or file.name in filename:
                    full_path = file
                    break
    
    if full_path.exists():
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(full_path), width=Inches(5.5))
            
            if caption:
                figure_counter += 1
                clean_caption = normalize_caption(caption, 'Figure', figure_counter)
                
                caption_para = doc.add_paragraph(clean_caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font_style(caption_para.runs[0], size=10, italic=True)
                set_paragraph_format(caption_para, space_before=3, space_after=3)
                
        except Exception as e:
            doc.add_paragraph(f"[Image non disponible : {Path(image_path).name}]")
    else:
        doc.add_paragraph(f"[Image non disponible : {Path(image_path).name}]")

def add_table(doc, headers, rows):
    """Ajouter un tableau avec style uniforme."""
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = clean_inline_text(header)
        set_font_style(header_cells[i].paragraphs[0].runs[0], size=11, bold=True)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = clean_inline_text(str(cell_data))
            set_font_style(row_cells[i].paragraphs[0].runs[0], size=11)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

def parse_markdown_table(table_text):
    """Parser un tableau Markdown"""
    lines = table_text.strip().split('\n')
    if len(lines) < 2:
        return [], []
    
    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    
    rows = []
    for line in lines[2:]:
        if line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    return headers, rows

def add_page_break(doc):
    """Ajouter un saut de page"""
    doc.add_page_break()

def add_toc(doc):
    """Ajouter une table des matières automatique"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Table des matières')
    run.bold = True
    run.font.size = Pt(16)
    set_font_style(run, size=16, bold=True)
    
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    set_paragraph_format(toc_para, space_before=6, space_after=12)

def add_figures_list(doc):
    """Ajouter la liste des figures"""
    global figures_list

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Liste des figures')
    run.bold = True
    set_font_style(run, size=16, bold=True)
    set_paragraph_format(paragraph, space_before=12, space_after=6)
    
    for fig in figures_list:
        para = doc.add_paragraph(fig)
        set_font_style(para.runs[0], size=12)
        set_paragraph_format(para, space_before=3, space_after=3)

def add_tables_list(doc):
    """Ajouter la liste des tableaux"""
    global tables_list

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Liste des tableaux')
    run.bold = True
    set_font_style(run, size=16, bold=True)
    set_paragraph_format(paragraph, space_before=12, space_after=6)
    
    for table in tables_list:
        para = doc.add_paragraph(table)
        set_font_style(para.runs[0], size=12)
        set_paragraph_format(para, space_before=3, space_after=3)

def add_annexes_list(doc):
    """Ajouter la liste des annexes"""
    global annexes_list

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Liste des annexes')
    run.bold = True
    set_font_style(run, size=16, bold=True)
    set_paragraph_format(paragraph, space_before=12, space_after=6)

    for annexe in annexes_list:
        para = doc.add_paragraph(annexe)
        set_font_style(para.runs[0], size=12)
        set_paragraph_format(para, space_before=3, space_after=3)

def convert_markdown_to_docx(md_file, output_file):
    """Convertir Markdown en Word avec style académique complet"""
    global figure_counter, table_counter, figures_list, tables_list, annexes_list
    
    # Réinitialiser les compteurs
    figure_counter = 0
    table_counter = 0
    figures_list = []
    tables_list = []
    annexes_list = []
    
    doc = Document()
    
    # Configuration du document
    set_margins(doc)
    set_header(doc)
    set_footer(doc)
    
    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    figures_list, tables_list, annexes_list = scan_document_lists(content)
    
    lines = content.split('\n')
    i = 0
    preliminary_lists_added = False
    
    # Traitement séquentiel complet
    while i < len(lines):
        line = lines[i]
        
        # Page de garde et résumé
        if i < 50 and line.startswith('#') and not re.match(r'^##\s+\d+\.', line.strip()):
            heading_text = line.lstrip('#').strip()
            if 'Page de garde' not in heading_text:
                add_heading(doc, heading_text, 0)
            i += 1
            continue
        
        # Lignes de la page de garde
        if i < 50 and line.startswith('**'):
            para = doc.add_paragraph(clean_inline_text(line))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font_style(para.runs[0], size=12, bold=True)
            i += 1
            continue
        
        # Sépar avant table des matières
        if line.strip() == '---' and i < 100:
            i += 1
            continue
        
        # Table des matières et listes après le résumé
        if not preliminary_lists_added and line.strip().startswith('## 1. Introduction'):
            add_page_break(doc)
            add_toc(doc)
            add_page_break(doc)
            add_figures_list(doc)
            add_page_break(doc)
            add_tables_list(doc)
            add_page_break(doc)
            add_annexes_list(doc)
            add_page_break(doc)
            preliminary_lists_added = True

        # Titres
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            if level <= 6:
                heading_text = line.lstrip('#').strip()
                if re.match(r'^Tableau\b', heading_text, flags=re.IGNORECASE):
                    table_counter += 1
                    caption_para = doc.add_paragraph(normalize_caption(heading_text, 'Tableau', table_counter))
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font_style(caption_para.runs[0], size=10, italic=True)
                    set_paragraph_format(caption_para, space_before=6, space_after=3)
                    i += 1
                    continue
                add_heading(doc, heading_text, max(level - 1, 0))
                i += 1
                continue
        
        # Séparateurs
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Blocs de code
        if line.strip().startswith('```'):
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            para = doc.add_paragraph(code_text)
            set_font_style(para.runs[0], font_name='Courier New', size=9)
            set_paragraph_format(para, line_spacing=1.0)
            i += 1
            continue
        
        # Images
        if '![' in line:
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                image_path = match.group(2)
                caption = None
                caption_index = i + 1
                while caption_index < len(lines) and not lines[caption_index].strip():
                    caption_index += 1
                if caption_index < len(lines) and lines[caption_index].strip().startswith('*Figure'):
                    caption = lines[caption_index].strip().strip('*')
                    i = caption_index
                add_image(doc, image_path, caption)
                i += 1
                continue
        
        # Tableaux
        if '|' in line and line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            
            table_text = '\n'.join(table_lines)
            headers, rows = parse_markdown_table(table_text)
            if headers and rows:
                add_table(doc, headers, rows)
            continue
        
        # Listes à puces
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                list_items.append(clean_inline_text(lines[i].strip()[2:]))
                i += 1
            for item in list_items:
                para = doc.add_paragraph(item, style='List Bullet')
                set_font_style(para.runs[0], size=12)
                set_paragraph_format(para)
            continue
        
        # Listes numérotées
        if re.match(r'^\d+\.', line.strip()):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.', lines[i].strip()):
                item = re.sub(r'^\d+\.\s*', '', lines[i].strip())
                list_items.append(clean_inline_text(item))
                i += 1
            for item in list_items:
                para = doc.add_paragraph(item, style='List Number')
                set_font_style(para.runs[0], size=12)
                set_paragraph_format(para)
            continue
        
        # Lignes vides
        if not line.strip():
            i += 1
            continue
        
        # Paragraphes normaux
        if line.strip():
            add_paragraph(doc, line.strip())
            i += 1
            continue
        
        i += 1
    
    # Sauvegarder
    doc.save(output_file)
    print(f"Document Word universitaire généré : {output_file}")
    print(f"\nStatistiques :")
    print(f"- Figures : {figure_counter}")
    print(f"- Tableaux : {table_counter}")

if __name__ == "__main__":
    print(f"Conversion académique de {MD_FILE} vers Word...")
    convert_markdown_to_docx(MD_FILE, WORD_OUTPUT)
    print("\nConversion terminée avec succès!")
    print("\nSpécifications appliquées :")
    print("- Police : Times New Roman")
    print("- Texte : 12 pt")
    print("- Titre principal : 18 pt Gras Centré")
    print("- Titres niveau 1 : 16 pt Gras")
    print("- Titres niveau 2 : 14 pt Gras")
    print("- Sous-sections : 13 pt")
    print("- Texte : Justifié")
    print("- Interligne : 1,5")
    print("- Espacement : 6 pt avant/après")
    print("- Marges : 2,5 cm")
    print("- En-tête : Rapport de projet tutoré – WattWatcher BF")
    print("- Pied de page : Numérotation")
    print("- Table des matières automatique")
    print("- Liste des figures")
    print("- Liste des tableaux")
    print("- Légendes des figures nettoyées")
    print("- Introductions et synthèses des chapitres")
